import os
import json
import random
import csv
import mimetypes
from collections import defaultdict
from datetime import datetime, timedelta
from typing import Dict, Any, Union, List, Optional
from pathlib import Path

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import google.generativeai as genai
except ImportError:
    raise ImportError("Install google-genai: pip install google-genai")

try:
    from dotenv import load_dotenv
except ImportError:
    raise ImportError("Install python-dotenv: pip install python-dotenv")

from agents.base_agent import BaseAgent, AgentResponse, AgentStatus

class ImageAgent(BaseAgent):
    """
    Image Agent - Produces quick visuals, mockups, and social-ready graphics from text prompts.
    
    This agent works in two steps:
    1. Generates a detailed JSON prompt from the user's topic
    2. Uses that prompt to generate an image with Gemini's image generation model
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(
            name="Image Agent",
            description="Produces quick visuals, mockups, and social-ready graphics from text prompts",
            config=config
        )
        
        # Load environment variables
        load_dotenv()
        
        self.api_key = self.config.get("api_key") or os.getenv("GEMINI_API_KEY")
        self.text_model_name = self.config.get("text_model_name", "gemini-2.0-flash")
        self.image_model_name = self.config.get("image_model_name", "gemini-3-pro-image-preview")
        self.save_to_file = self.config.get("save_to_file", True)
        self.output_dir = Path(self.config.get("output_dir", "outputs/images"))
        
        if self.save_to_file:
            self.output_dir.mkdir(parents=True, exist_ok=True)
        
        if self.api_key:
            genai.configure(api_key=self.api_key)
            self.text_model = genai.GenerativeModel(self.text_model_name)
            self.image_model = genai.GenerativeModel(self.image_model_name)
        else:
            self.text_model = None
            self.image_model = None
    
    def process(self, request: Dict[str, Any]) -> AgentResponse:
        """
        Process the image generation request.
        
        Args:
            request: Dictionary containing:
                - topic (required): The topic/description for the image
                - style (optional): Art style preference (e.g., "realistic", "cartoon", "minimalist")
                - aspect_ratio (optional): Image aspect ratio (e.g., "1:1", "16:9", "9:16")
                - filename (optional): Custom filename for the output
                
        Returns:
            AgentResponse with generated image path and metadata
        """
        self.status = AgentStatus.PROCESSING
        
        try:
            self.validate_request(request, ["topic"])
            
            topic = request["topic"]
            style = request.get("style", "professional and modern")
            aspect_ratio = request.get("aspect_ratio", "1:1")
            filename = request.get("filename")
            
            # Step 1: Generate detailed JSON prompt
            print(f"[IMAGE AGENT] Step 1: Generating detailed prompt from topic...")
            json_prompt = self._generate_detailed_prompt(topic, style, aspect_ratio)
            print(f"[IMAGE AGENT] Generated prompt: {json.dumps(json_prompt, indent=2)}")
            
            # Step 2: Generate image using the detailed prompt
            print(f"[IMAGE AGENT] Step 2: Generating image with Gemini...")
            image_path = self._generate_image(json_prompt, filename)
            print(f"[IMAGE AGENT] Image saved to: {image_path}")
            
            result = {
                "image_path": str(image_path),
                "prompt": json_prompt,
                "topic": topic,
                "style": style,
                "aspect_ratio": aspect_ratio,
            }
            
            self.status = AgentStatus.SUCCESS
            return AgentResponse(
                agent_name=self.name,
                status=AgentStatus.SUCCESS,
                result=result,
                metadata={
                    "topic": topic,
                    "prompt_length": len(json_prompt["detailed_description"]),
                    "model": self.image_model_name,
                }
            )
            
        except Exception as exc:
            self.status = AgentStatus.ERROR
            return AgentResponse(
                agent_name=self.name,
                status=AgentStatus.ERROR,
                result=None,
                error=str(exc),
                metadata={"request": request}
            )
    
    def _generate_detailed_prompt(self, topic: str, style: str, aspect_ratio: str) -> Dict[str, Any]:
        """
        Generate a detailed JSON prompt from the topic using the text model.
        
        Args:
            topic: User's topic/description
            style: Desired art style
            aspect_ratio: Image aspect ratio
            
        Returns:
            Dictionary with detailed prompt information
        """
        if not self.text_model:
            raise ValueError("Gemini API key required")
        
        system_prompt = """You are an expert image prompt engineer. Your job is to take a simple topic 
and expand it into a detailed, comprehensive image generation prompt.

Generate a JSON response with the following structure:
{
    "main_subject": "The primary subject of the image",
    "detailed_description": "A detailed, vivid description of what should be in the image (3-5 sentences)",
    "visual_elements": ["element1", "element2", "element3"],
    "composition": "How elements should be arranged (e.g., centered, rule of thirds)",
    "lighting": "Lighting description (e.g., soft natural light, dramatic studio lighting)",
    "color_palette": "Color scheme description (e.g., warm earth tones, vibrant and saturated)",
    "mood": "The emotional tone (e.g., professional, playful, serene)",
    "technical_details": "Camera/rendering details (e.g., 4K, sharp focus, depth of field)"
}

Make the prompt detailed, specific, and visually descriptive. Focus on what makes a great image."""

        user_prompt = f"""Topic: {topic}
Style: {style}
Aspect Ratio: {aspect_ratio}

Generate a detailed image prompt in JSON format. Be specific about visual details, composition, and atmosphere."""

        try:
            response = self.text_model.generate_content(f"{system_prompt}\\n\\n{user_prompt}")
            response_text = response.text.strip()
            
            # Extract JSON from response (handle code blocks)
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()
            
            prompt_data = json.loads(response_text)
            
            # Add metadata
            prompt_data["style"] = style
            prompt_data["aspect_ratio"] = aspect_ratio
            prompt_data["original_topic"] = topic
            
            return prompt_data
            
        except json.JSONDecodeError as e:
            print(f"[IMAGE AGENT] JSON parsing error: {e}")
            print(f"[IMAGE AGENT] Response text: {response_text}")
            # Fallback to a basic prompt structure
            return {
                "main_subject": topic,
                "detailed_description": f"A {style} image depicting {topic}",
                "visual_elements": [topic],
                "composition": "centered, balanced composition",
                "lighting": "natural, even lighting",
                "color_palette": "harmonious and appealing colors",
                "mood": "professional and polished",
                "technical_details": "high quality, sharp focus, 4K resolution",
                "style": style,
                "aspect_ratio": aspect_ratio,
                "original_topic": topic,
            }
        except Exception as e:
            print(f"[IMAGE AGENT] Error generating prompt: {e}")
            raise
    
    def _generate_image(self, json_prompt: Dict[str, Any], filename: Optional[str]) -> Path:
        """
        Generate an image using the detailed JSON prompt.
        
        Args:
            json_prompt: Detailed prompt dictionary
            filename: Optional custom filename
            
        Returns:
            Path to the saved image file
        """
        if not self.image_model:
            raise ValueError("Gemini image model not configured")
        
        # Construct the final image prompt from JSON
        image_prompt = self._construct_image_prompt(json_prompt)
        
        print(f"[IMAGE AGENT] Final prompt for image generation:")
        print(f"[IMAGE AGENT] {image_prompt}")
        
        try:
            # Generate image with Gemini
            response = self.image_model.generate_content(image_prompt)
            
            # Save the image
            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_topic = "".join(c for c in json_prompt["original_topic"][:30] if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_topic = safe_topic.replace(' ', '_')
                filename = f"image_{safe_topic}_{timestamp}"
            
            image_path = self.output_dir / f"{filename}.png"
            
            # Extract and save image from response
            if hasattr(response, '_result'):
                # Check if the response contains image data
                for candidate in response._result.candidates:
                    for part in candidate.content.parts:
                        if hasattr(part, 'inline_data'):
                            image_data = part.inline_data.data
                            image_path.write_bytes(image_data)
                            return image_path
            
            # If no image data found in response, try to extract it differently
            # This is a fallback in case the API response structure is different
            raise ValueError("No image data found in response. The model may not support image generation yet.")
            
        except Exception as e:
            print(f"[IMAGE AGENT] Error during image generation: {e}")
            print(f"[IMAGE AGENT] Note: Make sure you have access to the image generation model")
            raise
    
    def _construct_image_prompt(self, json_prompt: Dict[str, Any]) -> str:
        """
        Construct a comprehensive text prompt from the JSON structure.
        
        Args:
            json_prompt: Detailed prompt dictionary
            
        Returns:
            Formatted text prompt for image generation
        """
        parts = [
            f"Main Subject: {json_prompt.get('main_subject', '')}",
            f"\\nDescription: {json_prompt.get('detailed_description', '')}",
        ]
        
        if json_prompt.get('visual_elements'):
            elements = ", ".join(json_prompt['visual_elements'])
            parts.append(f"\\nVisual Elements: {elements}")
        
        parts.extend([
            f"\\nComposition: {json_prompt.get('composition', '')}",
            f"\\nLighting: {json_prompt.get('lighting', '')}",
            f"\\nColor Palette: {json_prompt.get('color_palette', '')}",
            f"\\nMood: {json_prompt.get('mood', '')}",
            f"\\nStyle: {json_prompt.get('style', '')}",
            f"\\nTechnical: {json_prompt.get('technical_details', '')}",
        ])
        
        return " ".join(parts)
    
    def get_capabilities(self) -> Dict[str, Any]:
        """Get agent capabilities and configuration."""
        base_caps = super().get_capabilities()
        base_caps.update({
            "text_model": self.text_model_name,
            "image_model": self.image_model_name,
            "has_api_key": bool(self.api_key),
            "output_directory": str(self.output_dir),
            "supported_styles": [
                "realistic",
                "cartoon",
                "minimalist",
                "watercolor",
                "professional",
                "modern",
                "vintage",
                "abstract",
            ],
            "supported_aspect_ratios": ["1:1", "16:9", "9:16", "4:3", "3:4"],
        })
        return base_caps

class DailyDigestSystem:
    def __init__(self):
        load_dotenv()
        self.api_key = os.getenv("GEMINI_API_KEY")
        if not self.api_key:
            print("Warning: GEMINI_API_KEY not found in environment variables. Please set it in .env file.")
        
        if self.api_key:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Initialize the Image Agent
        self.image_agent = ImageAgent()

    def get_data(self, file_paths: Union[str, List[str], None] = None) -> Dict[str, Any]:
        """
        Fetches data from multiple files (CSV, Excel, Docx, PDF, Image).
        Aggregates metrics by date across all files.
        If file_paths is None, prompts user for input.
        """
        # If no file_paths provided, prompt user
        if file_paths is None:
            print("\n[DIGEST SYSTEM] Enter file paths to process.")
            user_input = input("Enter file paths (comma-separated) or press Enter for default: ").strip()
            if user_input:
                file_paths = [p.strip() for p in user_input.split(",")]
            else:
                file_paths = ["data_digest_example_metrics.csv"]
        
        # Handle single string input
        if isinstance(file_paths, str):
            file_paths = [file_paths]
            
        # Filter out default if it doesn't exist and no other files provided
        valid_paths = []
        if file_paths:
            for p in file_paths:
                if os.path.exists(p):
                    valid_paths.append(p)
                else:
                    print(f"Warning: {p} not found.")
        
        if not valid_paths:
            print("No valid files found. Falling back to mock data.")
            return self._get_mock_data()

        aggregated_data = defaultdict(dict)
        all_dates = set()

        for file_path in valid_paths:
            ext = os.path.splitext(file_path)[1].lower()
            print(f"[DIGEST SYSTEM] Processing file: {file_path} ({ext})")
            
            file_data = {}
            if ext == '.csv':
                file_data = self._get_raw_data_from_csv(file_path)
            elif ext in ['.xlsx', '.xls']:
                file_data = self._get_raw_data_from_excel(file_path)
            elif ext == '.docx':
                struct = self._get_data_from_docx(file_path)
                file_data = self._unpack_struct(struct)
            elif ext == '.pdf':
                struct = self._get_data_from_pdf(file_path)
                file_data = self._unpack_struct(struct)
            elif ext in ['.png', '.jpg', '.jpeg', '.webp']:
                struct = self._get_data_from_image(file_path)
                file_data = self._unpack_struct(struct)
            else:
                print(f"Unsupported file extension: {ext}. Skipping.")
                continue
            
        if file_data:
                for date, metrics in file_data.items():
                    # Normalize dates to YYYY-MM-DD format for consistent comparison
                    normalized_date = self._normalize_date(date)
                    aggregated_data[normalized_date].update(metrics)
                    all_dates.add(normalized_date)
        
        if not all_dates:
            return self._get_mock_data()

        return self._process_dates(all_dates, aggregated_data)

    def _normalize_date(self, date_str: str) -> str:
        """Normalize various date formats to YYYY-MM-DD"""
        # If already in YYYY-MM-DD format, return as-is
        if isinstance(date_str, str) and len(date_str) == 10 and date_str[4] == '-' and date_str[7] == '-':
            return date_str
        
        # Otherwise, just return the string (Gemini will handle it)
        return str(date_str).split()[0]  # Remove timestamp if present

    def _unpack_struct(self, struct: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
        """Unpacks the yesterday/today structure back into date->metrics map"""
        data = {}
        if "yesterday" in struct and "date" in struct["yesterday"]:
            date = struct["yesterday"]["date"]
            metrics = struct["yesterday"].get("metrics", {})
            data[date] = metrics
        if "today" in struct and "date" in struct["today"]:
            date = struct["today"]["date"]
            metrics = struct["today"].get("metrics", {})
            data[date] = metrics
        return data

    def _get_raw_data_from_csv(self, file_path: str) -> Dict[str, Dict[str, Any]]:
        # Read CSV
        data_by_date = defaultdict(dict)

        try:
            with open(file_path, 'r', newline='') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    date = row.get('date')
                    metric = row.get('metric')
                    val_str = row.get('value')
                    
                    if not date or not metric or not val_str:
                        continue

                    try:
                        value = float(val_str)
                    except ValueError:
                        continue 
                    
                    data_by_date[date][metric] = value
        except Exception as e:
            print(f"Error reading CSV: {e}")
            return {}

        return data_by_date

    def _get_raw_data_from_excel(self, file_path: str) -> Dict[str, Dict[str, Any]]:
        if not pd:
            print("Pandas not installed.")
            return {}
        
        try:
            df = pd.read_excel(file_path)
            # Expect columns: date, metric, value
            if not {'date', 'metric', 'value'}.issubset(df.columns):
                 print("Excel missing required columns (date, metric, value).")
                 return {}
            
            data_by_date = defaultdict(dict)
            
            for _, row in df.iterrows():
                date = str(row['date']).split()[0] # Handle timestamp
                metric = row['metric']
                try:
                    value = float(row['value'])
                except (ValueError, TypeError):
                    continue
                
                data_by_date[date][metric] = value
                
            return data_by_date
            
        except Exception as e:
            print(f"Error reading Excel: {e}")
            return {}

    def _get_data_from_docx(self, file_path: str) -> Dict[str, Any]:
        if not docx:
            print("python-docx not installed.")
            return self._get_mock_data()
            
        try:
            doc = docx.Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    full_text += "\n" + " | ".join([cell.text for cell in row.cells])
            
            return self._extract_metrics_with_gemini(full_text, "text")
        except Exception as e:
            print(f"Error reading Docx: {e}")
            return self._get_mock_data()

    def _get_data_from_pdf(self, file_path: str) -> Dict[str, Any]:
        # Try using Gemini Multimodal first if file is small enough, 
        # otherwise extract text with pypdf
        
        # Option 1: Extract text
        if pypdf:
            try:
                reader = pypdf.PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return self._extract_metrics_with_gemini(text, "text")
            except Exception as e:
                print(f"Error reading PDF with pypdf: {e}")
        
        # Option 2: Pass file to Gemini (if we had a way to upload easily here without File API)
        # Since we are using genai.GenerativeModel, we can pass the file bytes if we use the File API
        # But for simplicity, let's stick to text extraction for PDF for now, 
        # or convert to image if we had pdf2image (which requires poppler).
        
        return self._get_mock_data()

    def _get_data_from_image(self, file_path: str) -> Dict[str, Any]:
        if not Image:
            print("Pillow not installed.")
            return self._get_mock_data()
            
        try:
            img = Image.open(file_path)
            return self._extract_metrics_with_gemini(img, "image")
        except Exception as e:
            print(f"Error reading Image: {e}")
            return self._get_mock_data()

    def _extract_metrics_with_gemini(self, content: Any, content_type: str) -> Dict[str, Any]:
        print(f"[DIGEST SYSTEM] Extracting metrics from {content_type} using Gemini...")
        
        prompt = """
        Extract business metrics from this content.
        Identify data for the two most recent dates available.
        
        Return ONLY a JSON object with this exact structure:
        {
            "yesterday": {
                "date": "YYYY-MM-DD",
                "metrics": { "metric_name": value, ... }
            },
            "today": {
                "date": "YYYY-MM-DD",
                "metrics": { "metric_name": value, ... }
            }
        }
        If exact dates aren't clear, infer "yesterday" and "today" as the two comparison points.
        """
        
        try:
            response = self.model.generate_content([prompt, content])
            text = response.text.strip()
            # Clean up markdown
            if "```json" in text:
                text = text.split("```json")[1].split("```")[0]
            elif "```" in text:
                text = text.split("```")[1].split("```")[0]
                
            return json.loads(text)
        except Exception as e:
            print(f"Error extracting with Gemini: {e}")
            return self._get_mock_data()

    def _detect_time_span(self, sorted_dates: List[str]) -> str:
        """
        Detect the time span covered by the data.
        Returns: "days", "weeks", "months", or "years"
        """
        if len(sorted_dates) < 2:
            return "days"
        
        try:
            first_date = datetime.strptime(sorted_dates[0], "%Y-%m-%d")
            last_date = datetime.strptime(sorted_dates[-1], "%Y-%m-%d")
            delta = (last_date - first_date).days
            
            if delta <= 7:
                return "days"
            elif delta <= 90:
                return "weeks"
            elif delta <= 730:
                return "months"
            else:
                return "years"
        except:
            return "days"

    def _process_dates(self, all_dates, data_by_date):
        """
        Process dates intelligently based on time span.
        Returns full historical data for analysis, not just 2 points.
        """
        sorted_dates = sorted(list(all_dates))
        
        if len(sorted_dates) < 2:
            print("Not enough data points to compare. Need at least 2 days.")
            return self._get_mock_data()

        time_span = self._detect_time_span(sorted_dates)
        
        return {
            "time_span": time_span,
            "data_points": len(sorted_dates),
            "date_range": {
                "start": sorted_dates[0],
                "end": sorted_dates[-1]
            },
            "historical_data": {date: data_by_date[date] for date in sorted_dates},
            "sorted_dates": sorted_dates
        }

    def _get_mock_data(self) -> Dict[str, Any]:
        """Fallback to mock data if CSV fails"""
        metrics = ["Active Users", "Revenue", "Server Load", "New Signups"]
        
        today = datetime.now().date()
        yesterday = today - timedelta(days=1)
        
        # Generate random data
        data = {
            "yesterday": {
                "date": str(yesterday),
                "metrics": {m: random.randint(100, 1000) for m in metrics}
            },
            "today": {
                "date": str(today),
                "metrics": {m: random.randint(100, 1000) for m in metrics}
            }
        }
        
        # Add some specific trends to make the summary interesting
        data["today"]["metrics"]["Revenue"] = int(data["yesterday"]["metrics"]["Revenue"] * 1.25)
        data["today"]["metrics"]["Server Load"] = int(data["yesterday"]["metrics"]["Server Load"] * 0.8)
        
        return data

    def generate_comparison_summary(self, data: Dict[str, Any]) -> str:
        """
        Uses Gemini to analyze the data intelligently.
        Handles both short-term comparisons (2 days) and long-term trends (weeks/months/years).
        Let's Gemini decide what the data is about.
        """
        if not self.api_key:
            return "Error: No API Key provided. Cannot generate summary."

        date_range = data.get('date_range', {})
        start_date = date_range.get('start', 'unknown')
        end_date = date_range.get('end', 'unknown')
        time_span = data.get('time_span', 'unknown')
        data_points = data.get('data_points', 'unknown')
        historical = json.dumps(data.get('historical_data', {}), indent=2)

        prompt = f"""
        You are an expert data analyst. Analyze the following dataset and provide insights.
        
        DATASET INFORMATION:
        - Time Span: {time_span}
        - Data Points: {data_points}
        - Date Range: {start_date} to {end_date}
        - Full Historical Data: {historical}
        
        YOUR TASK:
        1. **Identify what this data is about** - What business area/domain does it represent? (e.g., e-commerce, SaaS, operations, user engagement, etc.)
        2. **Understand the metrics** - What does each metric measure and why might it matter?
        3. **Analyze appropriately based on time span**:
           - If time_span is "days": Focus on day-to-day changes and immediate trends
           - If time_span is "weeks": Analyze week-over-week patterns and identify emerging trends
           - If time_span is "months": Look for seasonal patterns, growth trajectories, and major shifts
           - If time_span is "years": Analyze long-term growth, cycles, and strategic patterns
        4. **Calculate metrics**:
           - Growth rates (percentage change from start to end, and intermediate periods)
           - Trends: Is this metric going up, down, or stable?
           - Volatility: How consistent are the changes?
           - Peaks and valleys: When did highs and lows occur?
        5. **Generate insights**: What story does the data tell? What's working, what's declining, what's stable?
        6. **Write a comprehensive summary** (6-8 sentences) that:
           - Starts with what the data is about
           - Covers the overall trend and major patterns
           - Highlights significant changes and anomalies
           - Provides actionable insights
           - Is written in professional, clear language

        Remember: Do NOT assume specific metric names. Analyze whatever is present. Be flexible with the data structure.
        """
        
        try:
            response = self.model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            return f"Error generating summary: {str(e)}"

    def create_visual_report(self, summary: str, data: Dict[str, Any]) -> str:
        """
        Uses the ImageAgent to create a visual representation appropriate for the time span.
        For short periods: bar charts
        For longer periods: line charts showing trends
        """
        print("\n[DIGEST SYSTEM] Generating visual report...")
        
        time_span = data.get("time_span", "days")
        date_range = data.get("date_range", {})
        
        # Choose visualization type based on time span
        if time_span == "days":
            viz_type = "bar chart"
            description = "daily comparison"
        elif time_span == "weeks":
            viz_type = "line chart with trend"
            description = "weekly trend"
        elif time_span == "months":
            viz_type = "line chart with markers"
            description = "monthly progression"
        else:  # years
            viz_type = "line chart with area fill"
            description = "long-term trend analysis"
        
        request = {
            "topic": f"A professional {viz_type} visualization showing {description} from {date_range.get('start')} to {date_range.get('end')}. Metrics: {summary[:200]}...",
            "style": f"infographic, data visualization, {viz_type}, professional, clean, 2d vector art, business intelligence",
            "aspect_ratio": "16:9",
            "filename": f"report_{data.get('time_span', 'data')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        }
        
        response = self.image_agent.process(request)
        
        if response.status == AgentStatus.SUCCESS:
            return response.result["image_path"]
        else:
            raise Exception(f"Image generation failed: {response.error}")

    def run(self, file_paths: Union[str, List[str], None] = None):
        print("[DIGEST SYSTEM] Starting Intelligent Data Analysis...")
        
        # 1. Get Data
        data = self.get_data(file_paths)
        
        time_span = data.get("time_span", "unknown")
        data_points = data.get("data_points", "unknown")
        date_range = data.get("date_range", {})
        
        print(f"[DIGEST SYSTEM] Detected time span: {time_span}")
        print(f"[DIGEST SYSTEM] Data points: {data_points}")
        print(f"[DIGEST SYSTEM] Date range: {date_range.get('start')} to {date_range.get('end')}")
        print(f"[DIGEST SYSTEM] Raw Data: {json.dumps(data, indent=2)}")
        
        # 2. Generate Text Summary (Let Gemini decide what the data is about)
        print("\n[DIGEST SYSTEM] Analyzing data with Gemini...")
        summary = self.generate_comparison_summary(data)
        print("\n--- DATA ANALYSIS REPORT ---")
        print(summary)
        print("----------------------------\n")
        
        # 3. Generate Appropriate Visual
        if not summary.startswith("Error"):
            try:
                image_path = self.create_visual_report(summary, data)
                print(f"[DIGEST SYSTEM] Visual report generated: {image_path}")
                print(f"\n✓ Text Summary + Visual Report completed successfully!")
            except Exception as e:
                print(f"[DIGEST SYSTEM] Failed to generate visual report: {e}")
        else:
            print("[DIGEST SYSTEM] Skipping image generation due to analysis error.")

import argparse

if __name__ == "__main__":
    system = DailyDigestSystem()
    system.run()  # Will prompt for file paths
